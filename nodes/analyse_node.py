from groq import Groq
import docx
import json
import os
import base64
from dotenv import load_dotenv
load_dotenv()
GROQ_API_KEY = os.getenv("GROK_API_KEY")

client = Groq(api_key=GROQ_API_KEY)

def parse_srs_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def analyze_srs_text(srs_text):
    prompt = f"""
You are a senior backend architect. Analyze the given Software Requirements Specification (SRS) and extract the following technical elements in STRICT JSON format:


Return JSON like:
{{
  "api_endpoints": [
    {{
      "method": "GET",
      "path": "/users",
      "description": "Retrieve all users",
      "parameters": [{{ "name": "page", "type": "int", "required": false }}]
    }}
  ],
  "backend_logic": [
    "Users must be filtered by active status.",
    "If user role is admin, show all data."
  ],
  "database_schema": {{
    "tables": {{
      "users": {{
        "columns": {{
          "id": "integer",
          "name": "string",
          "email": "string",
          "created_at": "datetime"
        }},
        "primary_key": "id",
        "foreign_keys": []
      }}
    }}
  }},
  "authentication": {{
    "type": "JWT",
    "roles": ["admin", "user"],
    "rules": ["Admins can access all routes", "Users limited to /profile"]
  }}
}}


Be concise, structured, and correct. Avoid explanations. Only return the JSON.
----
SRS TEXT:
{srs_text}
"""
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return response.choices[0].message.content


def extract_images_from_docx(file_path):
    doc = docx.Document(file_path)
    image_bytes_list = []

    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            image_data = rel_obj.target_part.blob
            image_bytes_list.append(image_data)

    return image_bytes_list


def analyze_schema_from_image(image_bytes):
    encoded_image = base64.b64encode(image_bytes).decode("utf-8")
    vision_prompt = """
You are a backend engineer. Analyze this diagram and extract the database schema in JSON like:


{
  "database_schema": {
    "tables": {
      "table_name": {
        "columns": {
          "column_name": "data_type"
        },
        "primary_key": "column_name",
        "foreign_keys": [
          {"column": "name", "references": "other_table(column)"}
        ]
      }
    }
  }
}


Only respond with JSON. No markdown or extra text.
"""
    response = client.chat.completions.create(
        model="llama3-vision-70b",
        messages=[{
            "role": "user", "content": [
                {"type": "text", "text": vision_prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_image}"}}
            ]
        }],
        temperature=0.2
    )

    raw_content = response.choices[0].message.content.strip()
    try:
        json_start = raw_content.find("{")
        json_end = raw_content.rfind("}") + 1
        json_str = raw_content[json_start:json_end]
        return json.loads(json_str)
    except Exception as e:
        print("❌ Failed to parse JSON from vision response")
        print("Raw:", raw_content)
        return {}


def save_to_json(data, file_path="extracted_data.json"):
    if os.path.exists(file_path):
        with open(file_path, "r") as json_file:
            json_data = json.load(json_file)
    else:
        json_data = []

    json_data.append(data)

    with open(file_path, "w") as json_file:
        json.dump(json_data, json_file, indent=4)


def node1_parse_srs(state):
    file_path = state.get("srs_file", "srs.docx")
   
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return {"error": f"File not found: {file_path}"}
   
    srs_text = parse_srs_docx(file_path)
    result_str = analyze_srs_text(srs_text)

    print("\nRaw response from Groq:\n", result_str)

    try:
        if result_str.startswith('{"') and result_str.endswith('}'):
            result_str = result_str.strip()
        else:
            start_index = result_str.find("{")
            end_index = result_str.rfind("}") + 1
            result_str = result_str[start_index:end_index].strip()

        result_json = json.loads(result_str)
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        print("Raw response from Groq:", result_str)
        return {"error": str(e)}

    print("\n✅ Extracted System Requirements:\n")
    print(result_json)

    if not result_json.get("database_schema"):
        images = extract_images_from_docx(file_path)
        if images:
            print("🔍 No DB schema found in text. Trying to extract from image...")
            vision_result = analyze_schema_from_image(images[0])  
            if "database_schema" or "database" in vision_result:
                result_json["database_schema"] = vision_result["database_schema"]
                print("✅ Extracted DB schema from image.")

    save_to_json(result_json)

    return {
        "parsed_spec": result_json,
        "srs_text": srs_text
    }


